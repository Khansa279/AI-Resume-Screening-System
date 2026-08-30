"""Document parser for extracting text from PDF and DOCX files."""

import hashlib
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Literal
import re

# PDF parsing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

_CONTROL_CHAR_PATTERN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_PRIVATE_USE_PATTERN = re.compile(r"[\uE000-\uF8FF]")
_ZERO_WIDTH_PATTERN = re.compile(r"[\u200b\u200c\u200d\ufeff]")

# Line-leading bullet/list marker glyphs to normalize to "- " before any
# other sanitization runs. Deliberately anchored to the START of a line
# (optional leading whitespace only) and requires the marker to be
# followed by real content -- this is what keeps the normalization
# CONTEXTUAL: a hyphen, en dash, or arrow appearing inside a sentence,
# date, or range (e.g. "January 2022 - Present", "Jan 2022 - Dec 2023",
# "Q1 -> Q2", "co-founder") is never touched, because it isn't the first
# character on its line.
#
# Three buckets, chosen to cover "regardless of which bullet glyph or PDF
# font encoding was used" without guessing at every possible glyph:
#   1. Control characters (\x00-\x08, \x0b, \x0c, \x0e-\x1f, \x7f) -- PDF
#      icon fonts commonly render a bullet glyph as one of these when the
#      font has no real Unicode mapping for it. This is the confirmed
#      \x7f case (Priya Sharma, Rahul Verma, Ananya Patel, Vikram Singh's
#      resume PDFs), generalized to the whole control-char range instead
#      of hardcoding \x7f alone, since a different PDF export/font could
#      just as easily emit \x02 or \x0e for the same visual bullet.
#   2. Private Use Area characters (U+E000-U+F8FF) -- the other common
#      home for icon-font glyphs (Wingdings-style bullet fonts).
#   3. A curated set of standard Unicode bullet/list-marker symbols that
#      legitimately show up in resume PDFs: bullet, white bullet, small
#      black square, small triangle, triangular bullet, black circle,
#      black/white square, arrowhead bullets, middle dot, bullet
#      operator.
#   4. Plain "*" -- already recognized by ResumeParserAgent, included
#      here too only so it normalizes to the same canonical "- " marker.
#
# "-" itself is deliberately NOT in this set: it's already handled
# correctly downstream (ResumeParserAgent._parse_experience already
# recognizes a line-leading "-"), so leaving it alone avoids ANY risk of
# this new layer reinterpreting a line that happens to start with a
# hyphen (e.g. a standalone "- 2022" continuation fragment) differently
# than before. Em dash, en dash, and arrows are also deliberately
# excluded from the marker set for the same reason: those are common
# INSIDE dates/ranges, and only the confirmed, unambiguous bullet-glyph
# buckets above are normalized.
_BULLET_MARKER_LINE_PATTERN = re.compile(
    r"(?m)^([ \t]*)(?:"
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]"
    r"|[\uE000-\uF8FF]"
    r"|[\u2022\u25E6\u25AA\u25B8\u2023\u25CF\u25A0\u25A1\u27A4\u27A2\u00B7\u2219]"
    r"|\*"
    r")[ \t]*(?=\S)"
)


def normalize_bullet_markers(text: str) -> str:
    """Normalize line-leading bullet/list markers to the single canonical
    form ("- ") that ResumeParserAgent._parse_experience already
    recognizes, regardless of which bullet glyph or PDF font encoding
    produced them.

    Must run BEFORE _CONTROL_CHAR_PATTERN/_PRIVATE_USE_PATTERN strip
    those same characters to empty/space -- otherwise the one signal
    that distinguishes "this line is a bullet" from "this line is a new
    job title" is destroyed before this function ever sees it, which is
    exactly the bug this normalization fixes (a bullet glyph like \\x7f
    being deleted outright, leaving a bare line with no marker at all).

    Only touches the FIRST non-whitespace token of a line, and only when
    it's one of the specific bullet-glyph buckets in
    _BULLET_MARKER_LINE_PATTERN -- never touches punctuation appearing
    later in a line (dates, ranges, hyphenated words, arrows in prose).
    """
    if not text:
        return text
    return _BULLET_MARKER_LINE_PATTERN.sub(r"\1- ", text)


def normalize_extracted_text(text: str) -> str:
    """Collapse trivial whitespace so the same resume is hashed the same way
    even if PDF extractors differ on trailing spaces or extra blank lines."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    collapsed: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run == 1:
                collapsed.append("")
        else:
            blank_run = 0
            collapsed.append(line)
    return "\n".join(collapsed).strip()


def resume_content_hash(text: str) -> str:
    """SHA-256 of whitespace-normalized extracted resume text."""
    normalized = normalize_extracted_text(text)
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()

@dataclass
class ParseResult:
    """Result of document parsing."""
    text: str
    file_type: Literal["pdf", "docx", "txt", "unknown"]
    page_count: int = 1
    success: bool = True
    error_message: str = ""
    confidence: float = 1.0  # How confident we are in the extraction quality


class DocumentParser:
    """Parse PDF and DOCX documents to extract text content."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
    
    MIN_MEANINGFUL_CHARS = 40
    
    
    def __init__(self):
        """Initialize the document parser."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if required parsing libraries are available."""
        if not PDF_AVAILABLE:
            print("Warning: PyPDF2 not installed. PDF parsing will not work.")
        if not DOCX_AVAILABLE:
            print("Warning: python-docx not installed. DOCX parsing will not work.")
    
    def parse(self, file_path: str) -> ParseResult:
        """
        Parse a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParseResult with extracted text and metadata
        """
        path = Path(file_path)
        
        # Validate file exists
        if not path.exists():
            return ParseResult(
                text="",
                file_type="unknown",
                success=False,
                error_message=f"File not found: {file_path}",
                confidence=0.0
            )
        
        # Get file extension
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return ParseResult(
                text="",
                file_type="unknown",
                success=False,
                error_message=f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}",
                confidence=0.0
            )
        
        # Parse based on file type
        try:
            if ext == ".pdf":
                return self._parse_pdf(path)
            elif ext in {".docx", ".doc"}:
                return self._parse_docx(path)
            elif ext == ".txt":
                return self._parse_txt(path)
            else:
                return ParseResult(
                    text="",
                    file_type="unknown",
                    success=False,
                    error_message=f"No parser available for: {ext}",
                    confidence=0.0
                )
        except Exception as e:
            return ParseResult(
                text="",
                file_type=ext.replace(".", ""),  # type: ignore
                success=False,
                error_message=f"Error parsing file: {str(e)}",
                confidence=0.0
            )


    @staticmethod
    def _sanitize_extracted_text(text: str) -> str:
        """Remove extraction artifacts (control bytes, icon-font glyphs,
        zero-width chars) that were never legible text -- doesn't touch
        real words or wording.

        normalize_bullet_markers() runs FIRST, before the control-char /
        private-use stripping below, so a bullet glyph rendered as a
        control character or PUA icon-font character (see its docstring)
        gets converted to "- " while it's still recognizable as a
        line-leading marker, instead of being silently deleted by the
        steps that follow.
        """
        if not text:
            return text
        text = normalize_bullet_markers(text)
        text = _CONTROL_CHAR_PATTERN.sub("", text)
        text = _PRIVATE_USE_PATTERN.sub(" ", text)
        text = _ZERO_WIDTH_PATTERN.sub("", text)
        return text

    @staticmethod
    def _count_bullet_lines(text: str) -> int:
        """Count lines that are recognized bullet/list items AFTER
        _sanitize_extracted_text has already run (so this counts the
        canonical "- " marker normalize_bullet_markers() produces from
        any of the bullet-glyph buckets it recognizes -- see that
        function's docstring). This is a content-agnostic proxy for
        "how much of this PDF's list structure survived extraction",
        used only to compare two extractions of the SAME document
        against each other -- never to judge a document in isolation."""
        return sum(1 for line in text.split("\n") if line.strip().startswith("- "))

    def _parse_pdf(self, path: Path) -> ParseResult:
        if not PDF_AVAILABLE:
            return ParseResult(
                text="", file_type="pdf", success=False,
                error_message="PyPDF2 not installed. Run: pip install PyPDF2",
                confidence=0.0,
            )

        pypdf2_text, pypdf2_pages, _ = self._extract_with_pypdf2(path)
        pypdf2_text = self._sanitize_extracted_text(pypdf2_text)
        best_text, page_count = pypdf2_text, pypdf2_pages

        if PDFPLUMBER_AVAILABLE:
            plumber_text, plumber_pages, _ = self._extract_with_pdfplumber(path)
            plumber_text = self._sanitize_extracted_text(plumber_text)

            if self._alpha_char_count(best_text) < self.MIN_MEANINGFUL_CHARS:
                # Existing low-content fallback, unchanged: PyPDF2 got
                # almost nothing usable at all -- take pdfplumber if it
                # did meaningfully better on raw content.
                if self._alpha_char_count(plumber_text) > self._alpha_char_count(best_text):
                    best_text, page_count = plumber_text, (plumber_pages or page_count)
            else:
                # Both extractors produced substantial text. PyPDF2 and
                # pdfplumber use different internal text-positioning
                # strategies, so for some PDF templates one can silently
                # detach a bullet/list-marker glyph from the line it
                # belongs to while the other keeps it attached (see
                # normalize_bullet_markers' docstring for why that
                # matters downstream -- ResumeParserAgent/_parse_list_
                # section rely on the marker surviving to know where one
                # list item ends and the next begins; lose it and
                # multiple distinct entries can get silently merged into
                # one). Comparing recovered bullet-line counts is a
                # general, content-agnostic quality signal for "which
                # extraction preserved this document's list structure
                # better" -- not a check tied to any specific resume.
                # Verified against the existing sample_data/resumes/*.pdf
                # corpus: for those files PyPDF2 recovers dramatically
                # MORE bullets than pdfplumber (their bullets render as
                # icon-font glyphs PyPDF2 happens to keep in place), so
                # this rule correctly keeps PyPDF2 for them -- it only
                # switches extractor when pdfplumber demonstrably
                # recovers substantially more list structure without
                # losing a large fraction of the extracted content.
                pypdf2_bullets = self._count_bullet_lines(pypdf2_text)
                plumber_bullets = self._count_bullet_lines(plumber_text)
                plumber_alpha = self._alpha_char_count(plumber_text)
                pypdf2_alpha = self._alpha_char_count(pypdf2_text)
                if (
                    plumber_bullets >= pypdf2_bullets + 2
                    and plumber_alpha >= pypdf2_alpha * 0.7
                ):
                    best_text, page_count = plumber_text, (plumber_pages or page_count)

        if self._alpha_char_count(best_text) < self.MIN_MEANINGFUL_CHARS:
            tried = "PyPDF2" + (" and pdfplumber" if PDFPLUMBER_AVAILABLE else "")
            return ParseResult(
                text="", file_type="pdf", page_count=page_count, success=False,
                error_message=f"No meaningful text extracted (tried {tried}). Likely a scanned image PDF; OCR not currently implemented.",
                confidence=0.0,
            )

        return ParseResult(
            text=best_text, file_type="pdf", page_count=page_count,
            success=True, confidence=self._estimate_extraction_confidence(best_text),
        )

    def _extract_with_pypdf2(self, path: Path) -> tuple[str, int, str]:
        try:
            reader = PdfReader(str(path))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(t for t in pages if t), len(reader.pages), ""
        except Exception as e:
            return "", 0, str(e)

    def _extract_with_pdfplumber(self, path: Path) -> tuple[str, int, str]:
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                return "\n\n".join(t for t in pages if t), len(pdf.pages), ""
        except Exception as e:
            return "", 0, str(e)

    @staticmethod
    def _alpha_char_count(text: str) -> int:
        return sum(1 for c in text if c.isalpha())
    def _extract_with_pypdf2(self, path: Path) -> tuple[str, int, str]:
        """Returns (text, page_count, error_message)."""
        try:
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages), len(reader.pages), ""
        except Exception as e:
            return "", 0, str(e)

    def _extract_with_pdfplumber(self, path: Path) -> tuple[str, int, str]:
        """Returns (text, page_count, error_message)."""
        try:
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                page_count = len(pdf.pages)
            return "\n\n".join(pages), page_count, ""
        except Exception as e:
            return "", 0, str(e)

    @staticmethod
    def _alpha_char_count(text: str) -> int:
        """Count of alphabetic characters only -- used instead of
        len(text) so whitespace-only or punctuation-only 'extraction'
        from a scanned PDF isn't mistaken for real content."""
        return sum(1 for c in text if c.isalpha())
    def _parse_docx(self, path: Path) -> ParseResult:
        """Parse a DOCX file."""
        if not DOCX_AVAILABLE:
            return ParseResult(
                text="",
                file_type="docx",
                success=False,
                error_message="python-docx not installed. Run: pip install python-docx",
                confidence=0.0
            )
        
        try:
            doc = Document(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            full_text = "\n".join(paragraphs)
            confidence = self._estimate_extraction_confidence(full_text)
            
            return ParseResult(
                text=full_text,
                file_type="docx",
                page_count=1,  # DOCX doesn't have clear page boundaries
                success=True,
                confidence=confidence
            )
        except Exception as e:
            return ParseResult(
                text="",
                file_type="docx",
                success=False,
                error_message=f"DOCX parsing error: {str(e)}",
                confidence=0.0
            )
    
    def _parse_txt(self, path: Path) -> ParseResult:
        """Parse a plain text file."""
        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    text = path.read_text(encoding=encoding)
                    return ParseResult(
                        text=text,
                        file_type="txt",
                        page_count=1,
                        success=True,
                        confidence=1.0  # Plain text is always high confidence
                    )
                except UnicodeDecodeError:
                    continue
            
            return ParseResult(
                text="",
                file_type="txt",
                success=False,
                error_message="Could not decode file with any supported encoding",
                confidence=0.0
            )
        except Exception as e:
            return ParseResult(
                text="",
                file_type="txt",
                success=False,
                error_message=f"Text file reading error: {str(e)}",
                confidence=0.0
            )
    
    def _estimate_extraction_confidence(self, text: str) -> float:
        """
        Estimate how confident we are in the text extraction quality.
        
        Lower confidence when:
        - Text is very short
        - High ratio of special characters (possible OCR issues)
        - Missing expected resume sections
        """
        if not text or len(text) < 100:
            return 0.3
        
        # Check for reasonable text length
        word_count = len(text.split())
        if word_count < 50:
            return 0.5
        
        # Check for garbled text (high special character ratio)
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
        if len(text) > 0 and special_chars / len(text) > 0.3:
            return 0.6
        
        # Check for common resume keywords
        resume_keywords = ["experience", "education", "skills", "work", "job", "email", "phone"]
        keyword_matches = sum(1 for kw in resume_keywords if kw.lower() in text.lower())
        
        if keyword_matches >= 4:
            return 0.95
        elif keyword_matches >= 2:
            return 0.85
        else:
            return 0.7


# Singleton instance
_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get the global document parser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def parse_document(file_path: str) -> ParseResult:
    """Convenience function to parse a document."""
    return get_document_parser().parse(file_path)
